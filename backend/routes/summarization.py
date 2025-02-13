from flask import Blueprint, request, jsonify, send_file
import openai
import os
from pymongo import MongoClient
from dotenv import load_dotenv
import fitz  # PyMuPDF for PDFs
import docx  # python-docx for DOCs
from werkzeug.utils import secure_filename
import traceback
import jwt  # PyJWT for decoding JWT tokens
from bson import ObjectId
from datetime import datetime
from flask_jwt_extended import jwt_required, get_jwt_identity
from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas

# Load environment variables
load_dotenv()

# Set up MongoDB connection
mongo_url = os.getenv("MONGO_URI")
if not mongo_url:
    print("ERROR: MONGO_URI is not set properly in the .env file.")

client = MongoClient(mongo_url)
db = client["summarizer_db"]
summaries_collection = db["summaries"]

# Set OpenAI API Key
openai.api_key = os.getenv("OPENAI_API_KEY")
jwt_secret = os.getenv("JWT_SECRET_KEY")

# Blueprint for summarization routes
summarization_bp = Blueprint("summarization", __name__)

def extract_user_id():
    """Extracts user ID from JWT token in request headers"""
    token = request.headers.get("Authorization")
    
    if not token:
        print("No Authorization token found in request headers")
        return None

    try:
        decoded_token = jwt.decode(token, jwt_secret, algorithms=["HS256"])
        print("Decoded Token:", decoded_token)  # Debugging
        return decoded_token.get("sub", {}).get("id")  # Extract user ID from "sub"
    except jwt.ExpiredSignatureError:
        print("Token expired")
        return None
    except jwt.InvalidTokenError:
        print("Invalid token")
        return None


def extract_text_from_pdf(file):
    """Extracts text from PDF files"""
    try:
        pdf_document = fitz.open(stream=file.read(), filetype="pdf") 
        text = "".join([page.get_text("text") + "\n" for page in pdf_document])
        return text
    except Exception as e:
        print(f"PDF Extraction Error: {e}")
        return None

def extract_text_from_docx(docx_path):
    """Extracts text from DOCX files"""
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

# Route to summarize text
@summarization_bp.route("/summarize", methods=["POST"])
@jwt_required()
def summarize_text():
    # Extract user details from JWT
    user_details = get_jwt_identity()
    user_id = user_details.get("id")  # Get user ID from JWT token

    if not user_id:
        return jsonify({"error": "Unauthorized"}), 401

    data = request.json
    text = data.get("text", "")
    summary_length = data.get("summary_length", "medium")  # Default to medium
    summary_tone = data.get("summary_tone", "professional")  # Default to neutral

    if not text.strip():
        return jsonify({"summary":""}), 400

    # length prompts
    length_prompts = {
        "short": "Provide a brief summary (1-2 sentences).",
        "medium": "Provide a balanced summary (3-5 sentences).",
        "detailed": "Provide a detailed summary (multiple paragraphs)."
    }

    # tone prompts
    tone_prompts = {
        "professional": "Provide a clear and concise summary with a formal tone.",
        "casual": "Summarize in a simple and engaging way, casual manner",
        "academic": "Summarize in a well-structured, research-based manner with formal language."
    }
    prompt = f"{length_prompts.get(summary_length, length_prompts['medium'])} {tone_prompts.get(summary_tone, tone_prompts['professional'])}\n\n{text}"

    try:
        #Generate the summary using OpenAI's GPT-3.5 model
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        summary = response["choices"][0]["message"]["content"]

        #Generate a title for the summary
        title_prompt = f"Generate a short, meaningful title for the following summary: \n\n{summary}"
        title_response = openai.ChatCompletion.create(
            model = "gpt-3.5-turbo",
            messages = [{"role": "user", "content": title_prompt}]
        )
        generated_title = title_response["choices"][0]["message"]["content"].strip()

        # # Save summary in MongoDB with user ID and length
        # insert_result = summaries_collection.insert_one({
        #     "user_id": user_id,
        #     "original_text": text,
        #     "summary": summary,
        #     "summary_length": summary_length,
        #     "title": generated_title,
        #     "summary_tone": summary_tone
        # })

        return jsonify({
            "summary": summary,
            # "summary_id": str(insert_result.inserted_id),
            "title": generated_title,
            "summary_tone": summary_tone
            })

    except Exception as e:
        print(f"ERROR: {str(e)}")
        return jsonify({"error": str(e)}), 500

# Route to upload a file and summarize its content	
@summarization_bp.route("/upload", methods=["POST"])
def upload_file():
    user_id = extract_user_id()
    if not user_id:
        return jsonify({"error": "Unauthorized"}), 401

    try:
        if "file" not in request.files:
            return jsonify({"error": "No file part"}), 400

        file = request.files["file"]
        summary_length = request.form.get("summary_length", "medium")  # Default to medium
        summary_tone = request.form.get("summary_tone", "professional")  # Default to professional

        if file.filename == "":
            return jsonify({"error": "No selected file"}), 400

        filename = secure_filename(file.filename)

        if filename.endswith(".pdf"):
            text = extract_text_from_pdf(file)
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(file)
        else:
            return jsonify({"error": "Unsupported file format"}), 400

        
        length_prompts = {
            "short": "Provide a brief summary (1-2 sentences).",
            "medium": "Provide a balanced summary (3-5 sentences).",
            "detailed": "Provide a detailed summary (multiple paragraphs)."
        }

        
        tone_prompts = {
            "professional": "Provide a clear and concise summary with a formal tone.",
            "casual": "Summarize in a simple and engaging way, casual manner",
            "academic": "Summarize in a well-structured, research-based manner with formal language."
        }

        prompt = f"{length_prompts.get(summary_length, length_prompts['medium'])} {tone_prompts.get(summary_tone, tone_prompts['professional'])}\n\n{text}"

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        summary = response["choices"][0]["message"]["content"]

        # Generate a title based on the summary
        title_prompt = f"Generate a short, meaningful title for the following summary:\n\n{summary}"
        title_response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": title_prompt}]
        )
        generated_title = title_response["choices"][0]["message"]["content"].strip()

        # Save summary in MongoDB with user ID and selected length
        summaries_collection.insert_one({
            "user_id": user_id,
            "original_text": text,
            "summary": summary,
            "summary_length": summary_length,
            "title": generated_title,
            "summary_tone": summary_tone
        })

        return jsonify({
            "summary": summary, 
            "summary_length": summary_length,
            "title": generated_title,
            "sumamry_tone": summary_tone
            })

    except Exception as e:
        print(f"ERROR: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# Route to get summaries history
@summarization_bp.route("/SummariesHistory", methods=["GET"])
@jwt_required()
def get_user_summaries():
    # Extract user details from JWT
    user_details = get_jwt_identity()
    user_id = user_details.get("id")  # Get user ID from JWT token

    if not user_id:
        return jsonify({"error": "Unauthorized"}), 401

    try:
        # Find summaries belonging to the user
        summaries = list(summaries_collection.find({"user_id": user_id}))

        # Convert _id to string and ensure created_at exists
        formatted_summaries = []
        for summary in summaries:
            formatted_summaries.append({
                "_id": str(summary["_id"]),  # Convert ObjectId to string
                "summary_length": summary.get("summary_length", "N/A"),
                "title": summary.get("title", "N/A"), # Default if missing
                "summary_tone": summary.get("summary_tone", "N/A")
            })

        return jsonify({"summaries": formatted_summaries}), 200

    except Exception as e:
        print(f"ERROR: {str(e)}")
        return jsonify({"error": str(e)}), 500
    
# Route to get a specific summary
@summarization_bp.route("/summary/<summary_id>", methods=["GET"])
@jwt_required()
def get_summary(summary_id):
    # Extract user details from JWT
    user_details = get_jwt_identity()
    user_id = user_details.get("id")

    if not user_id:
        return jsonify({"error": "Unauthorized"}), 401
    
    summary = summaries_collection.find_one({"_id": ObjectId(summary_id), "user_id": user_id})

    if not summary:
        return jsonify({"error": "Summary not found"}), 404
    
    return jsonify({
        "_id": str(summary["_id"]),
        "original_text": summary.get("original_text", "N/A"),
        "summary": summary.get("summary", "N/A"),
        "summary_length": summary.get("summary_length", "N/A"),
        "title": summary.get("title", "N/A"),
        "summaryTone": summary.get("summary_tone", "N/A")
    })


# # Route to delete a specific summary
@summarization_bp.route("/summary/<summary_id>", methods=["DELETE"])
@jwt_required()
def delete_summary(summary_id):
    # Extract user details from JWT
    user_details = get_jwt_identity()
    user_id = user_details.get("id")

    if not user_id:
        return jsonify({"error": "Unauthorized"}), 401
    
    result = summaries_collection.delete_one({"_id": ObjectId(summary_id), "user_id": user_id})

    if result.deleted_count == 0:
        return jsonify({"error": "Summary not found or Unauthorized"}), 404

    return jsonify({"message": "Summary deleted successfully"}), 200

@summarization_bp.route("/summary/download/<summary_id>", methods=["GET"])
@jwt_required()
def download_summary(summary_id):
    format_type = request.args.get("format", "txt")

    # Extract user details from JWT
    user_details = get_jwt_identity()
    user_id = user_details.get("id")

    try:
        summary = summaries_collection.find_one({"_id": ObjectId(summary_id), "user_id": user_id})
    except Exception as e:
        return jsonify({"error": "Invalid summary ID"}), 400

    if not summary:
        return jsonify({"error": "Summary not found"}), 404

    title = summary.get("title", "Summary")
    content = summary.get("summary", "Summary content not found")

    if format_type == "txt":
        return generate_txt(title, content)
    elif format_type == "docx":
        return generate_docx(title, content)
    elif format_type == "pdf":
        return generate_pdf(title, content)
    else:
        return jsonify({"error": "Invalid format"}), 400


def generate_txt(title, content):
    """Generate and return a TXT file"""
    file_data = BytesIO()
    file_data.write(f"{title}\n\n{content}".encode("utf-8"))
    file_data.seek(0)

    return send_file(file_data, mimetype="text/plain", as_attachment=True, download_name=f"{title}.txt")


def generate_docx(title, content):
    """Generate and return a DOCX file"""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)

    file_data = BytesIO()
    doc.save(file_data)
    file_data.seek(0)

    return send_file(file_data, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name=f"{title}.docx")


def generate_pdf(title, content):
    """Generate and return a PDF file"""
    file_data = BytesIO()
    pdf = canvas.Canvas(file_data)
    pdf.setTitle(title)
    
    # Draw title
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(100, 750, title)
    
    # Draw content with proper spacing
    pdf.setFont("Helvetica", 12)
    y = 730  # Start below title
    for line in content.split("\n"):
        if y < 50:  # Prevent writing beyond the page limit
            pdf.showPage()  # Create a new page
            y = 750  # Reset position for the new page
        pdf.drawString(100, y, line)
        y -= 20  # Move text down

    pdf.save()
    file_data.seek(0)

    return send_file(file_data, mimetype="application/pdf", as_attachment=True, download_name=f"{title}.pdf")
